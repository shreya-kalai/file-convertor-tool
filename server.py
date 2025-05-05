import socket
from PIL import Image
from pdf2image import convert_from_path
from reportlab.pdfgen import canvas
from pdf2docx import Converter
import pytesseract
import docx
import os
from docx2pdf import convert as docx_to_pdf
import fitz  # PyMuPDF


def convert_pdf_to_docx(pdf_file):
    print("[SERVER] Trying standard PDF to DOCX conversion...")
    docx_file = pdf_file.replace(".pdf", ".docx")
    cv = Converter(pdf_file)
    cv.convert(docx_file)
    cv.close()
    return docx_file

def convert_scanned_pdf_to_docx(pdf_file):
    print("[SERVER] Falling back to OCR-based PDF to DOCX conversion...")
    pages = convert_from_path(pdf_file, 300)
    doc = docx.Document()
    for i, page in enumerate(pages):
        text = pytesseract.image_to_string(page)
        doc.add_paragraph(text)
        print(f"[SERVER] OCR processed page {i+1}")
    output = pdf_file.replace(".pdf", "_ocr.docx")
    doc.save(output)
    return output

def handle_client(client_socket):
    try:
        metadata = client_socket.recv(1024).decode()
        filename, target_format = metadata.split(",")
        safe_name = filename.replace(" ", "_").replace("(", "").replace(")", "")
        received_file = f"received_{safe_name}"
        client_socket.sendall(b"ACK")

        with open(received_file, "wb") as f:
            while True:
                data = client_socket.recv(4096)
                if not data:
                    break
                f.write(data)
        print(f"[SERVER] File received: {received_file}")

        converted_file = None

        # Image → PDF
        if filename.endswith((".jpg", ".jpeg", ".png")) and target_format == "pdf":
            print("[SERVER] Converting image to PDF...")
            img = Image.open(received_file)
            if img.mode != "RGB":
                img = img.convert("RGB")
            converted_file = f"{safe_name.split('.')[0]}_converted.pdf"
            img.save(converted_file, "PDF")

        # PDF → Image
        elif filename.endswith(".pdf") and target_format == "png":
            print("[SERVER] Converting PDF to image...")
            images = convert_from_path(received_file, 300)
            converted_file = f"{safe_name.split('.')[0]}_page1.png"
            images[0].save(converted_file, "PNG")
        #image <-> image
        elif filename.endswith((".jpg", ".jpeg", ".png", ".webp", ".bmp")) and target_format in ["jpg", "jpeg", "png", "webp", "bmp"]:
            print(f"[SERVER] Converting image to {target_format.upper()}...")
            img = Image.open(received_file)
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            converted_file = f"{safe_name.split('.')[0]}_converted.{target_format}"
            img.save(converted_file, target_format.upper())

        #image to docx
        elif filename.endswith((".jpg", ".jpeg", ".png")) and target_format == "docx":
            print("[SERVER] Embedding image in DOCX...")
            document = docx.Document()
            document.add_paragraph("Image inserted:")
            document.add_picture(received_file, width=docx.shared.Inches(4))
            converted_file = f"{safe_name.split('.')[0]}.docx"
            document.save(converted_file)
            
        #image to docx
        elif filename.endswith(".docx") and target_format in ["png", "jpg"]:
            print("[SERVER] Converting DOCX to image...")
            
            # Step 1: Convert DOCX → PDF
            from docx2pdf import convert as docx_to_pdf
            intermediate_pdf = f"{safe_name.split('.')[0]}_temp.pdf"
            docx_to_pdf(received_file, intermediate_pdf)

            # Step 2: Convert PDF → Image
            from pdf2image import convert_from_path
            images = convert_from_path(intermediate_pdf, 300)
            
            # Save first page (for now — extend later for all pages)
            converted_file = f"{safe_name.split('.')[0]}_page1.{target_format}"
            images[0].save(converted_file, target_format.upper())


        # TXT → PDF
        elif filename.endswith(".txt") and target_format == "pdf":
            print("[SERVER] Converting TXT to PDF...")
            converted_file = f"{safe_name.split('.')[0]}_converted.pdf"
            c = canvas.Canvas(converted_file)
            with open(received_file, "r") as f:
                lines = f.readlines()
                y = 800
                for line in lines:
                    c.drawString(100, y, line.strip())
                    y -= 20
            c.save()

        # PDF → DOCX
        elif filename.endswith(".pdf") and target_format == "docx":
            try:
                converted_file = convert_pdf_to_docx(received_file)
            except Exception as e:
                print(f"[SERVER] Standard PDF to DOCX failed: {e}")
                converted_file = convert_scanned_pdf_to_docx(received_file)
        #DOCX -> PDF
        elif filename.endswith(".docx") and target_format == "pdf":
            print("[SERVER] Converting DOCX to PDF...")
            converted_file = f"{safe_name.split('.')[0]}_converted.pdf"
            docx_to_pdf(received_file, converted_file)
        #PDF -> TXT
        elif filename.endswith(".pdf") and target_format == "txt":
            print("[SERVER] Extracting text from PDF...")
            converted_file = f"{safe_name.split('.')[0]}.txt"
            with fitz.open(received_file) as pdf:
                text = ""
                for page in pdf:
                    text += page.get_text()
            with open(converted_file, "w") as f:
                f.write(text)


        if converted_file and os.path.exists(converted_file):
            with open(converted_file, "rb") as f:
                while True:
                    chunk = f.read(4096)
                    if not chunk:
                        break
                    client_socket.sendall(chunk)
            print(f"[SERVER] Sent: {converted_file}")
        else:
            client_socket.sendall(b"Conversion failed or unsupported.")
    except Exception as e:
        print(f"[SERVER ERROR] {e}")
        client_socket.sendall(f"[ERROR] {e}".encode())
    finally:
        client_socket.close()

server = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
server.bind(("0.0.0.0", 9999))
server.listen(1)
print("[SERVER] Listening on port 9999...")

while True:
    client, _ = server.accept()
    handle_client(client)